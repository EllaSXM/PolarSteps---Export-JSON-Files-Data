
#%% Imports
import json
import os.path
from datetime import datetime
from docx import Document  # To Manipulate Word Documents
from collections import defaultdict
from docx.shared import RGBColor #To format the text with color


#%% Data Processing

#Load the relevant JSON files provided by PolarSteps
with open("locations.json") as f:
    location_data = json.load(f)
    
with open("trip.json") as f:
    trip_data = json.load(f)   
    

#All_steps is the first level list from trip_date that I will be extracting the journal entries from. 
print(type(trip_data['all_steps']))
all_steps = trip_data["all_steps"]

    
# Used List Comprehension instead of For Loops to extract data. Much simpler code. 
location = [all_steps[i]['location']["name"] for i in range(len(all_steps))]
location2 = [all_steps[i]['location']["detail"] for i in range(len(all_steps))]
location3 = [all_steps[i]['location']["full_detail"] for i in range(len(all_steps))]
latitude = [all_steps[i]['location']["lat"] for i in range(len(all_steps))]
longitude = [all_steps[i]['location']["lon"] for i in range(len(all_steps))]   
country_code = [all_steps[i]['location']["country_code"] for i in range(len(all_steps))]
journal_text = [all_steps[i]['description'] for i in range(len(all_steps))]
journal_title = [all_steps[i]['display_name'] for i in range(len(all_steps))]
weather_desc = [all_steps[i]['weather_condition'] for i in range(len(all_steps))]
temperature = [all_steps[i]['weather_temperature'] for i in range(len(all_steps))]
date1 = [datetime.fromtimestamp(all_steps[i]['creation_time']).isoformat() for i in range(len(all_steps))]
date2 = [datetime.fromtimestamp(all_steps[i]['start_time']).isoformat() for i in range(len(all_steps))]

#Note if you receive an error "type object 'datetime has no attribute 'datetime' then run this version of the previous two lines instead:
#date1 = [datetime.datetime.fromtimestamp(all_steps[i]['creation_time']).isoformat() for i in range(len(all_steps))]
#date2 = [datetime.daterime.fromtimestamp(all_steps[i]['start_time']).isoformat() for i in range(len(all_steps))]


#Format Date
date1_formatted = [datetime.fromisoformat(date).strftime("%Y-%b-%d") for date in date1]



# This is the data I want to extract and save as my journal entries. 
data = {
    'date_formatted': date1_formatted,
    'Location': location,
    'Location 2': location2,
    'Title': journal_title,
    'Text': journal_text
}

#The data will be extracted in three ways. 
#1) All entries in one file.
#2) Entries separated by month.
#3) Each individual post will be extracted into a text file of its own. 


#%% Save journal entries into one file. 
def create_journal_all(data, filename):
    #Open an empty document
    doc = Document()
    
    for date, title, loc, loc2, text in zip(data['date_formatted'], data['Title'], data['Location'], data['Location 2'], data['Text']):
        # Create and format date paragraph
        p = doc.add_paragraph()
        date_run = p.add_run(date)
        date_run.bold = True

        # Create and format title paragraph
        p = doc.add_paragraph()
        title_run = p.add_run(title)
        title_run.font.color.rgb = RGBColor(0, 171, 126)

        # Create and format location and detail paragraph
        p = doc.add_paragraph()
        p.add_run(f'{loc}\n{loc2}')

        # Create and format text paragraph
        p = doc.add_paragraph(text)
    
    doc.save(filename)
    
        
#Run the function        
create_journal_all(data, 'MyTravelJournal.docx') 


#%% Save journal entries by month. 

def create_journal_by_month(data, filename_prefix):
    
    # Group entries by month and year
    monthly_data = defaultdict(list)
    for i, date in enumerate(data['date_formatted']):
        month_year = datetime.strptime(date, "%Y-%b-%d").strftime("%Y-%m")
        entry = {
            'date': data['date_formatted'][i],
            'location': data['Location'][i],
            'location2': data['Location 2'][i],
            'title': data['Title'][i],
            'text': data['Text'][i]
        }
        monthly_data[month_year].append(entry)
    
    # Create a document for each month
    for month_year, entries in monthly_data.items():
        doc = Document()
        for entry in entries:
            # Create and format date paragraph
            p = doc.add_paragraph()
            date_run = p.add_run(entry["date"])
            date_run.bold = True

            # Create and format title paragraph
            p = doc.add_paragraph()
            title_run = p.add_run(entry["title"])
            title_run.font.color.rgb = RGBColor(0, 171, 126)

            # Create and format location and detail paragraph
            p = doc.add_paragraph()
            p.add_run(f'{entry["location"]}\n{entry["location2"]}')

            # Create and format text paragraph
            p = doc.add_paragraph(entry["text"])
        filename = f'{filename_prefix}_{month_year}.docx'
        doc.save(filename)

# Create monthly journals
create_journal_by_month(data, 'journal')



#%% Save journal entries per day as text files. 

# Function to create individual text files
def create_individual_text_files(data):
    if not os.path.exists("journal_entries"):
        os.makedirs("journal_entries")
    for i, (date, title, loc, loc2, text) in enumerate(zip(data['date_formatted'], data['Title'], data['Location'], data['Location 2'], data['Text'])):
        filename = f"journal_entries/entry_{i+1}_{date}.txt"
        with open(filename, 'w') as file:
            file.write(f'Date: {date}\nTitle: {title}\nLocation: {loc}\nDetail: {loc2}\nText: {text}\n')

# Run the function to create individual text files
create_individual_text_files(data)


